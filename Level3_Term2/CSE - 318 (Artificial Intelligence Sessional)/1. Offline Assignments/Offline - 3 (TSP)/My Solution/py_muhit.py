import pandas as pd
from docx import Document
from tabulate import tabulate  # For better table formatting in console output

# Load the data
file_path = "Report.csv"  # Replace with the actual file path
data = pd.read_csv(file_path)

# Check the columns in the data
print(data.columns)  # To check if 'Tour_Cost' exists

# Strip any leading or trailing spaces in the column names
data.columns = data.columns.str.strip()

# Extract methods
constructive_methods = ['Nearest_Neighbour', 'Nearest_Insertion', 'Cheapest_Insertion']
perturbative_methods = ['2-opt', 'Node_Swap', 'Node_Shift']

# Calculate total experiments (assuming all rows are part of the experiment)
total_experiments = 21

# Filter data for constructive and perturbative methods
constructive_data = data[data['Construction_Method'].isin(constructive_methods)]
perturbative_data = data[data['Perturbative_Method'].isin(perturbative_methods)]

# Calculate average costs for constructive methods
constructive_avg_costs = (
    constructive_data.groupby('Construction_Method')['Tour_Cost']
    .mean()
    .reset_index()
    .rename(columns={'Tour_Cost': 'Average Cost'})
)

# Calculate average costs for perturbative methods
perturbative_avg_costs = (
    perturbative_data.groupby('Perturbative_Method')['Tour_Cost']
    .mean()
    .reset_index()
    .rename(columns={'Tour_Cost': 'Average Cost'})
)

# Save the results in a Word document
doc = Document()
doc.add_heading("TSP Analysis: Average Costs for Each Method", level=1)

# Add constructive methods table
doc.add_heading("Constructive Methods", level=2)
doc.add_paragraph("The table below shows the average costs for constructive methods.")
constructive_table = doc.add_table(rows=1, cols=2)
constructive_table.style = 'Table Grid'

# Add header
constructive_table.cell(0, 0).text = "Construction Heuristic"
constructive_table.cell(0, 1).text = "Average Cost"

# Add data rows
for _, row in constructive_avg_costs.iterrows():
    cells = constructive_table.add_row().cells
    cells[0].text = row['Construction_Method']
    cells[1].text = f"{row['Average Cost']:.2f}"

# Add perturbative methods table
doc.add_heading("Perturbative Methods", level=2)
doc.add_paragraph("The table below shows the average costs for perturbative methods.")
perturbative_table = doc.add_table(rows=1, cols=2)
perturbative_table.style = 'Table Grid'

# Add header
perturbative_table.cell(0, 0).text = "Improvement Heuristic"
perturbative_table.cell(0, 1).text = "Average Cost"

# Add data rows
for _, row in perturbative_avg_costs.iterrows():
    cells = perturbative_table.add_row().cells
    cells[0].text = row['Perturbative_Method']
    cells[1].text = f"{row['Average Cost']:.2f}"

# Save the document
doc.save("TSP_Average_Costs.docx")

print("Document created successfully: TSP_Average_Costs.docx")

# Print results to console for reference
print("\nConstructive Methods Average Costs:")
print(tabulate(constructive_avg_costs, headers="keys", tablefmt="pretty", showindex=False))

print("\nPerturbative Methods Average Costs:")
print(tabulate(perturbative_avg_costs, headers="keys", tablefmt="pretty", showindex=False))
