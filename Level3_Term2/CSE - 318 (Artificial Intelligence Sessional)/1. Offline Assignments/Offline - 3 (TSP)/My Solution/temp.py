import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import numpy as np  # For NaN handling
from docx.shared import Inches


# Load the data
data_text = """
File,Construction_Method,Time,Perturbative_Method,Time, Tour_Cost
berlin52.tsp,Nearest_Neighbour,0,None,0,8980.92
berlin52.tsp,Nearest_Neighbour,0,Node_Swap,31,8667.84
berlin52.tsp,Nearest_Neighbour,0,Node_Shift,141,8320.34
berlin52.tsp,Nearest_Neighbour,0,2-opt,0,8384.19
berlin52.tsp,Cheapest_Insertion,0,None,0,9014.89
berlin52.tsp,Cheapest_Insertion,0,Node_Swap,16,8889.92
berlin52.tsp,Cheapest_Insertion,0,Node_Shift,79,8400.46
berlin52.tsp,Cheapest_Insertion,0,2-opt,0,8545.14
berlin52.tsp,Cheapest_Insertion,0,None,0,9014.89
berlin52.tsp,Cheapest_Insertion,0,Node_Swap,15,8889.92
berlin52.tsp,Cheapest_Insertion,0,Node_Shift,79,8400.46
berlin52.tsp,Cheapest_Insertion,0,2-opt,0,8545.14
bier127.tsp,Nearest_Neighbour,0,None,0,135752
bier127.tsp,Nearest_Neighbour,0,Node_Swap,550,128780
bier127.tsp,Nearest_Neighbour,0,Node_Shift,1775,124006
bier127.tsp,Nearest_Neighbour,0,2-opt,15,126605
bier127.tsp,Cheapest_Insertion,64,None,0,140691
bier127.tsp,Cheapest_Insertion,64,Node_Swap,393,138984
bier127.tsp,Cheapest_Insertion,64,Node_Shift,1269,133194
bier127.tsp,Cheapest_Insertion,64,2-opt,0,134889
bier127.tsp,Cheapest_Insertion,79,None,0,140691
bier127.tsp,Cheapest_Insertion,79,Node_Swap,377,138984
bier127.tsp,Cheapest_Insertion,79,Node_Shift,1286,133194
bier127.tsp,Cheapest_Insertion,79,2-opt,16,134889
ch130.tsp,Nearest_Neighbour,0,None,0,7575.29
ch130.tsp,Nearest_Neighbour,0,Node_Swap,552,7404.23
ch130.tsp,Nearest_Neighbour,0,Node_Shift,1429,6971.08
ch130.tsp,Nearest_Neighbour,0,2-opt,0,7025.87
ch130.tsp,Cheapest_Insertion,63,None,0,7279.21
ch130.tsp,Cheapest_Insertion,63,Node_Swap,423,6882.32
ch130.tsp,Cheapest_Insertion,63,Node_Shift,2032,6748
ch130.tsp,Cheapest_Insertion,63,2-opt,16,6685.54
ch130.tsp,Cheapest_Insertion,188,None,0,7279.21
ch130.tsp,Cheapest_Insertion,188,Node_Swap,1270,6882.32
ch130.tsp,Cheapest_Insertion,188,Node_Shift,2633,6748
ch130.tsp,Cheapest_Insertion,188,2-opt,17,6685.54
eil51.tsp,Nearest_Neighbour,0,None,0,513.61
eil51.tsp,Nearest_Neighbour,0,Node_Swap,140,478.869
eil51.tsp,Nearest_Neighbour,0,Node_Shift,188,448.579
eil51.tsp,Nearest_Neighbour,0,2-opt,16,465.909
eil51.tsp,Cheapest_Insertion,0,None,0,496.252
eil51.tsp,Cheapest_Insertion,0,Node_Swap,78,480.249
eil51.tsp,Cheapest_Insertion,0,Node_Shift,188,456.015
eil51.tsp,Cheapest_Insertion,0,2-opt,15,457.737
eil51.tsp,Cheapest_Insertion,0,None,0,496.252
eil51.tsp,Cheapest_Insertion,0,Node_Swap,80,480.249
eil51.tsp,Cheapest_Insertion,0,Node_Shift,189,456.015
eil51.tsp,Cheapest_Insertion,0,2-opt,3,457.737
eil76.tsp,Nearest_Neighbour,0,None,0,711.993
eil76.tsp,Nearest_Neighbour,0,Node_Swap,338,634.699
eil76.tsp,Nearest_Neighbour,0,Node_Shift,660,575.742
eil76.tsp,Nearest_Neighbour,0,2-opt,0,620.03
eil76.tsp,Cheapest_Insertion,31,None,0,612.391
eil76.tsp,Cheapest_Insertion,31,Node_Swap,252,597.695
eil76.tsp,Cheapest_Insertion,31,Node_Shift,659,585.949
eil76.tsp,Cheapest_Insertion,31,2-opt,16,597.695
eil76.tsp,Cheapest_Insertion,32,None,0,612.391
eil76.tsp,Cheapest_Insertion,32,Node_Swap,251,597.695
eil76.tsp,Cheapest_Insertion,32,Node_Shift,644,585.949
eil76.tsp,Cheapest_Insertion,32,2-opt,0,597.695
eil101.tsp,Nearest_Neighbour,0,None,0,825.242
eil101.tsp,Nearest_Neighbour,0,Node_Swap,754,742.836
eil101.tsp,Nearest_Neighbour,0,Node_Shift,1490,727.216
eil101.tsp,Nearest_Neighbour,0,2-opt,15,707.523
eil101.tsp,Cheapest_Insertion,78,None,0,702.96
eil101.tsp,Cheapest_Insertion,78,Node_Swap,393,689.193
eil101.tsp,Cheapest_Insertion,78,Node_Shift,1145,668.047
eil101.tsp,Cheapest_Insertion,78,2-opt,0,689.193
eil101.tsp,Cheapest_Insertion,93,None,0,702.96
eil101.tsp,Cheapest_Insertion,93,Node_Swap,377,689.193
eil101.tsp,Cheapest_Insertion,93,Node_Shift,1179,668.047
eil101.tsp,Cheapest_Insertion,93,2-opt,0,689.193
kroA100.tsp,Nearest_Neighbour,0,None,0,26856.4
kroA100.tsp,Nearest_Neighbour,0,Node_Swap,740,24574.7
kroA100.tsp,Nearest_Neighbour,0,Node_Shift,1572,23647.6
kroA100.tsp,Nearest_Neighbour,0,2-opt,6,23641.4
kroA100.tsp,Cheapest_Insertion,80,None,0,24307.8
kroA100.tsp,Cheapest_Insertion,80,Node_Swap,547,23556.6
kroA100.tsp,Cheapest_Insertion,80,Node_Shift,1180,22692.1
kroA100.tsp,Cheapest_Insertion,80,2-opt,0,23556.6
kroA100.tsp,Cheapest_Insertion,95,None,0,24307.8
kroA100.tsp,Cheapest_Insertion,95,Node_Swap,534,23556.6
kroA100.tsp,Cheapest_Insertion,95,Node_Shift,1200,22692.1
kroA100.tsp,Cheapest_Insertion,95,2-opt,14,23556.6
kroB100.tsp,Nearest_Neighbour,5,None,0,29155
kroB100.tsp,Nearest_Neighbour,5,Node_Swap,609,28594.2
kroB100.tsp,Nearest_Neighbour,5,Node_Shift,3130,25242.9
kroB100.tsp,Nearest_Neighbour,5,2-opt,15,24985.2
kroB100.tsp,Cheapest_Insertion,79,None,0,25580.9
kroB100.tsp,Cheapest_Insertion,79,Node_Swap,393,24704.5
kroB100.tsp,Cheapest_Insertion,79,Node_Shift,1528,24082.6
kroB100.tsp,Cheapest_Insertion,79,2-opt,15,24374.8
kroB100.tsp,Cheapest_Insertion,79,None,0,25580.9
kroB100.tsp,Cheapest_Insertion,79,Node_Swap,393,24704.5
kroB100.tsp,Cheapest_Insertion,79,Node_Shift,1491,24082.6
kroB100.tsp,Cheapest_Insertion,79,2-opt,1,24374.8
kroC100.tsp,Nearest_Neighbour,16,None,0,26327.4
kroC100.tsp,Nearest_Neighbour,16,Node_Swap,580,25713.2
kroC100.tsp,Nearest_Neighbour,16,Node_Shift,1558,24044.9
kroC100.tsp,Nearest_Neighbour,16,2-opt,16,23665.6
kroC100.tsp,Cheapest_Insertion,94,None,0,25262.2
kroC100.tsp,Cheapest_Insertion,94,Node_Swap,596,24447.6
kroC100.tsp,Cheapest_Insertion,94,Node_Shift,2352,22911.5
kroC100.tsp,Cheapest_Insertion,94,2-opt,16,24341.7
kroC100.tsp,Cheapest_Insertion,94,None,0,25262.2
kroC100.tsp,Cheapest_Insertion,94,Node_Swap,580,24447.6
kroC100.tsp,Cheapest_Insertion,94,Node_Shift,2323,22911.5
kroC100.tsp,Cheapest_Insertion,94,2-opt,15,24341.7
kroD100.tsp,Nearest_Neighbour,0,None,0,26950.5
kroD100.tsp,Nearest_Neighbour,0,Node_Swap,723,26580.8
kroD100.tsp,Nearest_Neighbour,0,Node_Shift,2341,24468.8
kroD100.tsp,Nearest_Neighbour,0,2-opt,0,24727.1
kroD100.tsp,Cheapest_Insertion,95,None,0,25204.3
kroD100.tsp,Cheapest_Insertion,95,Node_Swap,567,24967.1
kroD100.tsp,Cheapest_Insertion,95,Node_Shift,1175,24089.3
kroD100.tsp,Cheapest_Insertion,95,2-opt,10,24251.7
kroD100.tsp,Cheapest_Insertion,79,None,0,25204.3
kroD100.tsp,Cheapest_Insertion,79,Node_Swap,584,24967.1
kroD100.tsp,Cheapest_Insertion,79,Node_Shift,1131,24089.3
kroD100.tsp,Cheapest_Insertion,79,2-opt,15,24251.7
kroE100.tsp,Nearest_Neighbour,0,None,0,27587.2
kroE100.tsp,Nearest_Neighbour,0,Node_Swap,768,26081.4
kroE100.tsp,Nearest_Neighbour,0,Node_Shift,1521,25422.1
kroE100.tsp,Nearest_Neighbour,0,2-opt,16,24216
kroE100.tsp,Cheapest_Insertion,93,None,0,25902
kroE100.tsp,Cheapest_Insertion,93,Node_Swap,595,25065.3
kroE100.tsp,Cheapest_Insertion,93,Node_Shift,1534,23164.3
kroE100.tsp,Cheapest_Insertion,93,2-opt,0,23338
kroE100.tsp,Cheapest_Insertion,94,None,0,25902
kroE100.tsp,Cheapest_Insertion,94,Node_Swap,580,25065.3
kroE100.tsp,Cheapest_Insertion,94,Node_Shift,1569,23164.3
kroE100.tsp,Cheapest_Insertion,94,2-opt,0,23338
lin105.tsp,Nearest_Neighbour,0,None,0,20362.8
lin105.tsp,Nearest_Neighbour,0,Node_Swap,660,19631.4
lin105.tsp,Nearest_Neighbour,0,Node_Shift,2309,18379.3
lin105.tsp,Nearest_Neighbour,0,2-opt,16,18538.4
lin105.tsp,Cheapest_Insertion,110,None,0,16934.6
lin105.tsp,Cheapest_Insertion,110,Node_Swap,676,16538.7
lin105.tsp,Cheapest_Insertion,110,Node_Shift,1573,16000.7
lin105.tsp,Cheapest_Insertion,110,2-opt,14,15849.7
lin105.tsp,Cheapest_Insertion,50,None,0,16934.6
lin105.tsp,Cheapest_Insertion,50,Node_Swap,489,16538.7
lin105.tsp,Cheapest_Insertion,50,Node_Shift,1857,16000.7
lin105.tsp,Cheapest_Insertion,50,2-opt,16,15849.7
pr76.tsp,Nearest_Neighbour,0,None,0,153462
pr76.tsp,Nearest_Neighbour,0,Node_Swap,267,123434
pr76.tsp,Nearest_Neighbour,0,Node_Shift,1084,133923
pr76.tsp,Nearest_Neighbour,0,2-opt,15,126350
pr76.tsp,Cheapest_Insertion,39,None,0,125936
pr76.tsp,Cheapest_Insertion,39,Node_Swap,356,120443
pr76.tsp,Cheapest_Insertion,39,Node_Shift,737,117799
pr76.tsp,Cheapest_Insertion,39,2-opt,17,119413
pr76.tsp,Cheapest_Insertion,30,None,0,125936
pr76.tsp,Cheapest_Insertion,30,Node_Swap,331,120443
pr76.tsp,Cheapest_Insertion,30,Node_Shift,799,117799
pr76.tsp,Cheapest_Insertion,30,2-opt,0,119413
pr124.tsp,Nearest_Neighbour,0,None,0,69299.4
pr124.tsp,Nearest_Neighbour,0,Node_Swap,1476,65186.5
pr124.tsp,Nearest_Neighbour,0,Node_Shift,2209,64381.9
pr124.tsp,Nearest_Neighbour,0,2-opt,15,61910.5
pr124.tsp,Cheapest_Insertion,158,None,0,65318.2
pr124.tsp,Cheapest_Insertion,158,Node_Swap,1130,63935.6
pr124.tsp,Cheapest_Insertion,158,Node_Shift,1472,61440.1
pr124.tsp,Cheapest_Insertion,158,2-opt,16,62237.5
pr124.tsp,Cheapest_Insertion,172,None,0,65318.2
pr124.tsp,Cheapest_Insertion,172,Node_Swap,1067,63935.6
pr124.tsp,Cheapest_Insertion,172,Node_Shift,1459,61440.1
pr124.tsp,Cheapest_Insertion,172,2-opt,15,62237.5
pr144.tsp,Nearest_Neighbour,0,None,0,61650.7
pr144.tsp,Nearest_Neighbour,0,Node_Swap,1146,61419.9
pr144.tsp,Nearest_Neighbour,0,Node_Shift,5826,61012.9
pr144.tsp,Nearest_Neighbour,0,2-opt,15,61399.2
pr144.tsp,Cheapest_Insertion,222,None,0,73033.1
pr144.tsp,Cheapest_Insertion,222,Node_Swap,1199,72562.9
pr144.tsp,Cheapest_Insertion,222,Node_Shift,5857,72116.4
pr144.tsp,Cheapest_Insertion,222,2-opt,32,59680
pr144.tsp,Cheapest_Insertion,267,None,0,73033.1
pr144.tsp,Cheapest_Insertion,267,Node_Swap,1136,72562.9
pr144.tsp,Cheapest_Insertion,267,Node_Shift,5833,72116.4
pr144.tsp,Cheapest_Insertion,267,2-opt,32,59680
rat99.tsp,Nearest_Neighbour,0,None,0,1564.72
rat99.tsp,Nearest_Neighbour,0,Node_Swap,1052,1416.68
rat99.tsp,Nearest_Neighbour,0,Node_Shift,1879,1362.67
rat99.tsp,Nearest_Neighbour,0,2-opt,17,1378.44
rat99.tsp,Cheapest_Insertion,79,None,0,1482.02
rat99.tsp,Cheapest_Insertion,79,Node_Swap,877,1412.23
rat99.tsp,Cheapest_Insertion,79,Node_Shift,1586,1279.47
rat99.tsp,Cheapest_Insertion,79,2-opt,16,1364.87
rat99.tsp,Cheapest_Insertion,78,None,0,1482.02
rat99.tsp,Cheapest_Insertion,78,Node_Swap,923,1412.23
rat99.tsp,Cheapest_Insertion,78,Node_Shift,1539,1279.47
rat99.tsp,Cheapest_Insertion,78,2-opt,16,1364.87
st70.tsp,Nearest_Neighbour,0,None,0,805.531
st70.tsp,Nearest_Neighbour,0,Node_Swap,125,776.279
st70.tsp,Nearest_Neighbour,0,Node_Shift,392,750.851
st70.tsp,Nearest_Neighbour,0,2-opt,0,751.273
st70.tsp,Cheapest_Insertion,31,None,0,778.995
st70.tsp,Cheapest_Insertion,31,Node_Swap,188,742.434
st70.tsp,Cheapest_Insertion,31,Node_Shift,376,737.916
st70.tsp,Cheapest_Insertion,31,2-opt,0,726.333
st70.tsp,Cheapest_Insertion,16,None,0,778.995
st70.tsp,Cheapest_Insertion,16,Node_Swap,188,742.434
st70.tsp,Cheapest_Insertion,16,Node_Shift,363,737.916
st70.tsp,Cheapest_Insertion,16,2-opt,15,726.333
ch150.tsp,Nearest_Neighbour,0,None,0,8194.61
ch150.tsp,Nearest_Neighbour,0,Node_Swap,1316,8119.93
ch150.tsp,Nearest_Neighbour,0,Node_Shift,12066,7328.93
ch150.tsp,Nearest_Neighbour,0,2-opt,46,7194.1
ch150.tsp,Cheapest_Insertion,284,None,0,7994.29
ch150.tsp,Cheapest_Insertion,284,Node_Swap,2037,7696.01
ch150.tsp,Cheapest_Insertion,284,Node_Shift,9369,7002.79
ch150.tsp,Cheapest_Insertion,284,2-opt,17,7466.9
ch150.tsp,Cheapest_Insertion,283,None,0,7994.29
ch150.tsp,Cheapest_Insertion,283,Node_Swap,2019,7696.01
ch150.tsp,Cheapest_Insertion,283,Node_Shift,9257,7002.79
ch150.tsp,Cheapest_Insertion,283,2-opt,32,7466.9
lin318.tsp,Nearest_Neighbour,21,None,0,54033.6
lin318.tsp,Nearest_Neighbour,21,Node_Swap,25254,53239.4
lin318.tsp,Nearest_Neighbour,21,Node_Shift,90204,49528.9
lin318.tsp,Nearest_Neighbour,21,2-opt,205,49639.4
lin318.tsp,Cheapest_Insertion,2712,None,0,49454.8
lin318.tsp,Cheapest_Insertion,2712,Node_Swap,19532,48321.2
lin318.tsp,Cheapest_Insertion,2712,Node_Shift,52768,46302
lin318.tsp,Cheapest_Insertion,2712,2-opt,126,46068
lin318.tsp,Cheapest_Insertion,2775,None,0,49454.8
lin318.tsp,Cheapest_Insertion,2775,Node_Swap,19270,48321.2
lin318.tsp,Cheapest_Insertion,2775,Node_Shift,53187,46302
lin318.tsp,Cheapest_Insertion,2775,2-opt,126,46068
a280.tsp,Nearest_Neighbour,0,None,0,3148.11
a280.tsp,Nearest_Neighbour,0,Node_Swap,8971,3100.49
a280.tsp,Nearest_Neighbour,0,Node_Shift,62622,2893.54
a280.tsp,Nearest_Neighbour,0,2-opt,141,2797.03
a280.tsp,Cheapest_Insertion,1801,None,0,3101.79
a280.tsp,Cheapest_Insertion,1801,Node_Swap,13108,3040.45
a280.tsp,Cheapest_Insertion,1801,Node_Shift,35696,2784.08
a280.tsp,Cheapest_Insertion,1801,2-opt,94,2826.15
a280.tsp,Cheapest_Insertion,1912,None,0,3101.79
a280.tsp,Cheapest_Insertion,1912,Node_Swap,13238,3040.45
a280.tsp,Cheapest_Insertion,1912,Node_Shift,35573,2784.08
a280.tsp,Cheapest_Insertion,1912,2-opt,111,2826.15
rat195.tsp,Nearest_Neighbour,0,None,0,2761.96
rat195.tsp,Nearest_Neighbour,0,Node_Swap,8421,2690.39
rat195.tsp,Nearest_Neighbour,0,Node_Shift,17006,2528.38
rat195.tsp,Nearest_Neighbour,0,2-opt,78,2499.61
rat195.tsp,Cheapest_Insertion,629,None,0,2814.57
rat195.tsp,Cheapest_Insertion,629,Node_Swap,4370,2691.87
rat195.tsp,Cheapest_Insertion,629,Node_Shift,14315,2507.72
rat195.tsp,Cheapest_Insertion,629,2-opt,31,2593
rat195.tsp,Cheapest_Insertion,661,None,0,2814.57
rat195.tsp,Cheapest_Insertion,661,Node_Swap,4349,2691.87
rat195.tsp,Cheapest_Insertion,661,Node_Shift,14265,2507.72
rat195.tsp,Cheapest_Insertion,661,2-opt,31,2593

"""  # Example: Replace with your full dataset

from io import StringIO
data = pd.read_csv(StringIO(data_text.strip()))

# Clean column names by stripping leading/trailing spaces
data.columns = data.columns.str.strip()

# Replace 'None' with NaN in the Perturbative_Method column
data['Perturbative_Method'] = data['Perturbative_Method'].replace('None', np.nan)

# Convert relevant columns to numeric for analysis
data['Construction_Time'] = pd.to_numeric(data['Time'])
data['Perturbative_Time'] = pd.to_numeric(data['Perturbative_Method'], errors='coerce')  # coerce invalid values to NaN
data['Tour_Cost'] = pd.to_numeric(data['Tour_Cost'], errors='coerce')

# Create a summary table
summary = data.groupby(['Construction_Method', 'Perturbative_Method']).agg(
    Avg_Tour_Cost=('Tour_Cost', 'mean'),
    Avg_Total_Time=('Perturbative_Time', 'mean')
).reset_index()

# Plotting
plt.figure(figsize=(12, 6))
sns.barplot(data=summary, x="Construction_Method", y="Avg_Tour_Cost", hue="Perturbative_Method")
plt.title("Average Tour Cost by Method Combination")
plt.ylabel("Average Tour Cost")
plt.xlabel("Construction Method")
plt.legend(title="Perturbative Method")
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('Tour_Cost_Barplot.png')

plt.figure(figsize=(12, 6))
sns.barplot(data=summary, x="Construction_Method", y="Avg_Total_Time", hue="Perturbative_Method")
plt.title("Average Time by Method Combination")
plt.ylabel("Average Time (ms)")
plt.xlabel("Construction Method")
plt.legend(title="Perturbative Method")
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('Time_Barplot.png')

# Creating the document
doc = Document()
doc.add_heading('TSP Algorithm Performance Report', level=1)

# Add text description
doc.add_paragraph(
    "This report analyzes the performance of different combinations of constructive "
    "and perturbative methods for solving the Travelling Salesman Problem (TSP). "
    "The analysis is based on two metrics: Tour Cost and Time Taken."
)

# Add the summary table
doc.add_heading('Summary Table', level=2)
table = doc.add_table(rows=1, cols=len(summary.columns))
table.style = 'Light List'
hdr_cells = table.rows[0].cells
for i, col_name in enumerate(summary.columns):
    hdr_cells[i].text = col_name

for _, row in summary.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

# Add graphs
doc.add_heading('Graphs', level=2)

doc.add_paragraph("Figure 1: Average Tour Cost by Method Combination.")
doc.add_picture('Tour_Cost_Barplot.png', width=Inches(5))

doc.add_paragraph("Figure 2: Average Time by Method Combination.")
doc.add_picture('Time_Barplot.png', width=Inches(5))

# Save the report
doc.save('TSP_Performance_Report.docx')

print("Report generated: TSP_Performance_Report.docx")
