import os
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# List of metrics and their corresponding file suffixes
metrics = {
    'Congestion Window': 'cwnd',
    'Slow Start Threshold': 'ssth',
    'Round Trip Time': 'rtt',
    'Retransmission Timeout': 'rto',
    'Next TX Sequence Number': 'next-tx',
    'In-flight Bytes': 'inflight',
    'Next RX Sequence Number': 'next-rx'
}

flows = range(7)  # Flows from flow0 to flow6
colors = ['orange', 'lime', 'blue', 'magenta', 'cyan', 'yellow', 'red']  # Colors for different flows

# Create a Word document
document = Document()
document.add_heading('TCP HTCP Modified Metrics Analysis', 0)

# Create a directory for the plots
os.makedirs('plots', exist_ok=True)

for metric_name, suffix in metrics.items():
    plt.figure(figsize=(10, 6))
    for flow in flows:
        filename = f'TcpVariantsComparison-flow{flow}-{suffix}.data'
        if not os.path.isfile(filename):
            continue
        # Load data
        data = np.loadtxt(filename)
        if data.size == 0:
            continue
        time = data[:, 0]
        value = data[:, 1]
        plt.plot(time, value, label=f'Flow {flow}', color=colors[flow % len(colors)])

    plt.title(f'{metric_name} over Time')
    plt.xlabel('Time (s)')
    plt.ylabel(metric_name)
    plt.legend()
    plt.grid(True)
    # Save the plot
    plot_filename = f'plots/{suffix}_plot.png'
    plt.savefig(plot_filename)
    plt.close()

    # Add the plot to the Word document
    document.add_heading(metric_name, level=1)
    document.add_picture(plot_filename, width=Inches(6))
    document.add_page_break()

# Save the Word document
document.save('TcpHtcpModified_Metrics.docx')
