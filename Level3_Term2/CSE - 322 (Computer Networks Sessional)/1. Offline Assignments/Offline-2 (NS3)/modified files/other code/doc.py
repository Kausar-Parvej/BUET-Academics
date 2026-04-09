import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

# Read the CSV files (ensure these files are in the same directory as the script)
aodv_data = pd.read_csv("aodv.csv")
raodv_data = pd.read_csv("raodv.csv")

# Rename columns to align with the plotting requirements
aodv_data.rename(columns={
    "numNodes": "NumNodes",
    "speed": "NodeSpeed",
    "packetRate": "PacketsPerSecond",
    "Throughput (kbps)": "Throughput",
    "End-to-End Delay (s)": "Delay",
    "PDR (%)": "DeliveryRatio",
    "Packet Drop Ratio (%)": "DropRatio"
}, inplace=True)

raodv_data.rename(columns={
    "numNodes": "NumNodes",
    "speed": "NodeSpeed",
    "packetRate": "PacketsPerSecond",
    "Throughput (kbps)": "Throughput",
    "End-to-End Delay (s)": "Delay",
    "PDR (%)": "DeliveryRatio",
    "Packet Drop Ratio (%)": "DropRatio"
}, inplace=True)

# Define specific slices for each category
aodv_data_node = aodv_data.iloc[0:4]  # Top 4 entries (NumNodes varying)
raodv_data_node = raodv_data.iloc[0:4]

aodv_data_speed = aodv_data.iloc[4:8]  # Middle 4 entries (NodeSpeed varying)
raodv_data_speed = raodv_data.iloc[4:8]

aodv_data_pps = aodv_data.iloc[8:12]  # Last 4 entries (PacketsPerSecond varying)
raodv_data_pps = raodv_data.iloc[8:12]

# Define metrics and their corresponding data slices
metrics = [
    # For NumNodes
    (aodv_data_node, raodv_data_node, "NumNodes", "Throughput", "Node Number", "Throughput (kbps)", "Node Number vs Throughput", "node_vs_throughput.png"),
    (aodv_data_node, raodv_data_node, "NumNodes", "Delay", "Node Number", "Delay (s)", "Node Number vs Delay", "node_vs_delay.png"),
    (aodv_data_node, raodv_data_node, "NumNodes", "DeliveryRatio", "Node Number", "Delivery Ratio (%)", "Node Number vs Delivery Ratio", "node_vs_delivery_ratio.png"),
    (aodv_data_node, raodv_data_node, "NumNodes", "DropRatio", "Node Number", "Drop Ratio (%)", "Node Number vs Drop Ratio", "node_vs_drop_ratio.png"),

    # For NodeSpeed
    (aodv_data_speed, raodv_data_speed, "NodeSpeed", "Throughput", "Node Speed (m/s)", "Throughput (kbps)", "Speed vs Throughput", "speed_vs_throughput.png"),
    (aodv_data_speed, raodv_data_speed, "NodeSpeed", "Delay", "Node Speed (m/s)", "Delay (s)", "Speed vs Delay", "speed_vs_delay.png"),
    (aodv_data_speed, raodv_data_speed, "NodeSpeed", "DeliveryRatio", "Node Speed (m/s)", "Delivery Ratio (%)", "Speed vs Delivery Ratio", "speed_vs_delivery_ratio.png"),
    (aodv_data_speed, raodv_data_speed, "NodeSpeed", "DropRatio", "Node Speed (m/s)", "Drop Ratio (%)", "Speed vs Drop Ratio", "speed_vs_drop_ratio.png"),

    # For PacketsPerSecond
    (aodv_data_pps, raodv_data_pps, "PacketsPerSecond", "Throughput", "Packets Per Second", "Throughput (kbps)", "Packets Per Second vs Throughput", "pps_vs_throughput.png"),
    (aodv_data_pps, raodv_data_pps, "PacketsPerSecond", "Delay", "Packets Per Second", "Delay (s)", "Packets Per Second vs Delay", "pps_vs_delay.png"),
    (aodv_data_pps, raodv_data_pps, "PacketsPerSecond", "DeliveryRatio", "Packets Per Second", "Delivery Ratio (%)", "Packets Per Second vs Delivery Ratio", "pps_vs_delivery_ratio.png"),
    (aodv_data_pps, raodv_data_pps, "PacketsPerSecond", "DropRatio", "Packets Per Second", "Drop Ratio (%)", "Packets Per Second vs Drop Ratio", "pps_vs_drop_ratio.png"),
]

# Function to generate plots and save them
def generate_plot(x, y_aodv, y_raodv, xlabel, ylabel, title, file_name):
    plt.figure(figsize=(6, 4))
    plt.plot(x, y_aodv, label="AODV", color='blue', marker='o', linestyle='-')
    plt.plot(x, y_raodv, label="RAODV", color='red', marker='x', linestyle='--')
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.title(title)
    plt.legend()
    plt.grid(True)
    plt.savefig(file_name)
    plt.close()

# Create the Word document
doc = Document()
doc.add_heading('Network Performance Comparison (AODV vs RAODV)', 0)

# Generate plots and add them to the Word document
for data_aodv, data_raodv, x_col, y_col, xlabel, ylabel, title, file_name in metrics:
    x_aodv = data_aodv[x_col]
    y_aodv = data_aodv[y_col]
    x_raodv = data_raodv[x_col]
    y_raodv = data_raodv[y_col]
    generate_plot(x_aodv, y_aodv, y_raodv, xlabel, ylabel, title, file_name)
    doc.add_paragraph(f'{title}:')
    doc.add_picture(file_name, width=Inches(5.0))

# Save the Word document
doc.save("network_performance_comparison.docx")

# Clean up the images
for file_name in os.listdir('.'):
    if file_name.endswith('.png'):
        os.remove(file_name)

print("Word document has been generated: network_performance_comparison.docx")
